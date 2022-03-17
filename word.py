import re
from docx import Document

# file = Document("./words/2330_108年第四季中文版合併財報.docx")

# all_text = ""
# for index, context in enumerate(file.paragraphs):
#     all_text+=context.text

def split_pages(text: str)-> tuple:
    """
    Use regular expression to find "page", like "- 12 -"
    return: a tuple about splited page content and splited page numbers
    """
    page_content = re.split("-\s*[0-9]+\s*-", text)[:-1]
    page_num = re.findall("-\s*[0-9]+\s*-", text)
    page_num = [num[1:-2] for num in page_num ]
    return (page_num, page_content)

def re_keywords_string(keywords:str)->str:
    keywords = list(keywords)
    regular_keywords = ""
    for index, keyword in enumerate(keywords):
        if index != len(keywords)-1:
            regular_keywords = regular_keywords + keyword + "\s*"
        else: regular_keywords = regular_keywords + keyword
    
    return regular_keywords

def find_keywords_pages(re_keywords: str, splited_pages: list, splited_num:list)-> list:
    """
    Find keyword in page content(str), and return page number
    return: page number where keywords occured.
    """
    re_keywords = re.compile(re_keywords)
    in_ = []
    for index, page_content in enumerate(splited_pages):
        if re.search(re_keywords, page_content) != None :
            in_.append(splited_num[index])
    return in_


# x= "aa aa - 1 - a  aa - 3 - aa   a - 5 - a   aa - 7 - bb - 12 - bb - 12 - bbbbbb - 12 - bbbbbb - 12 - cccc - 123 - cccc - 123 - cccc - 123 - cccc - 123 - cccc - 123 - cccc - 123 -"

# nums, contents = split_page(all_text)
# re_keywords = re_keywords_string("外匯")
# print(find_keywords_page(re_keywords, contents, nums))

# tables = file.tables
# for table in tables